import streamlit as st
import os
import tempfile
import docx
import json
import re
import difflib
import matplotlib.pyplot as plt
import numpy as np
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from tabulate import tabulate
import nltk
import base64
from io import BytesIO
from collections import Counter

# Download NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)

# Set page configuration
st.set_page_config(
    page_title="QAJ Format Checker by Ridwan Marqas",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Add custom CSS
st.markdown("""
<style>
    .main {
        padding: 2rem;
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: 2px;
    }
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        white-space: pre-wrap;
        background-color: #f0f2f6;
        border-radius: 4px 4px 0 0;
        gap: 1px;
        padding-top: 10px;
        padding-bottom: 10px;
    }
    .stTabs [aria-selected="true"] {
        background-color: #e6f0ff;
    }
    .compliance-compliant {
        color: green;
        font-weight: bold;
    }
    .compliance-non-compliant {
        color: red;
        font-weight: bold;
    }
    .compliance-partially-compliant {
        color: orange;
        font-weight: bold;
    }
    h1 {
        margin-bottom: 1.5rem;
    }
    h2 {
        margin-top: 2rem;
        margin-bottom: 1rem;
    }
    h3 {
        margin-top: 1.5rem;
        margin-bottom: 0.75rem;
    }
    .highlight {
        background-color: #ffffcc;
        padding: 0.2rem;
        border-radius: 0.2rem;
    }
    .table-container {
        overflow-x: auto;
    }
    .citation-example {
        font-family: monospace;
        background-color: #f5f5f5;
        padding: 0.2rem 0.4rem;
        border-radius: 0.2rem;
    }
    .critical-issue {
        color: #721c24;
        background-color: #f8d7da;
        padding: 0.75rem;
        margin-bottom: 1rem;
        border: 1px solid #f5c6cb;
        border-radius: 0.25rem;
    }
</style>
""", unsafe_allow_html=True)

class DocumentAnalyzer:
    """Class to analyze and compare DOCX documents for format compliance."""
    
    def __init__(self, template_file, article_file):
        """Initialize with template and article file objects."""
        self.template_file = template_file
        self.article_file = article_file
        
        # Create temporary directory
        self.temp_dir = tempfile.mkdtemp()
        
        # Initialize data structures
        self.template_info = None
        self.article_info = None
        self.comparison_results = {
            "structure": [],
            "formatting": [],
            "content": [],
            "citations": [],
            "tables_figures": [],
            "references": [],
            "section_numbering": [],
            "academic_quality": [],
            "journal_citation": [],
            "self_citation": [],
            "major_issues": []
        }
    
    def extract_document_info(self, file_obj, output_prefix):
        """Extract document information including styles, formatting, and structure."""
        # Save file to temporary location
        temp_path = os.path.join(self.temp_dir, f"{output_prefix}.docx")
        with open(temp_path, 'wb') as f:
            f.write(file_obj.getvalue())
        
        # Open document
        doc = docx.Document(temp_path)
        
        # Document information
        doc_info = {
            "filename": file_obj.name,
            "paragraphs_count": len(doc.paragraphs),
            "sections_count": len(doc.sections),
            "tables_count": len(doc.tables),
            "styles": {},
            "paragraphs": [],
            "sections": [],
            "tables": [],
            "figures": [],
            "citations": [],
            "references": [],
            "authors": [],
            "affiliations": []
        }
        
        # Extract styles information
        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                style_info = {
                    "name": style.name,
                    "type": "paragraph",
                    "font": None,
                    "font_size": None,
                    "bold": None,
                    "italic": None,
                    "alignment": None,
                    "spacing_before": None,
                    "spacing_after": None,
                    "line_spacing": None
                }
                
                if style.font:
                    style_info["font"] = style.font.name if style.font.name else None
                    style_info["font_size"] = style.font.size.pt if style.font.size else None
                    style_info["bold"] = style.font.bold if hasattr(style.font, 'bold') else None
                    style_info["italic"] = style.font.italic if hasattr(style.font, 'italic') else None
                
                if style.paragraph_format:
                    pf = style.paragraph_format
                    style_info["alignment"] = str(pf.alignment) if pf.alignment else None
                    style_info["spacing_before"] = pf.space_before.pt if pf.space_before else None
                    style_info["spacing_after"] = pf.space_after.pt if pf.space_after else None
                    style_info["line_spacing"] = pf.line_spacing if pf.line_spacing else None
                
                doc_info["styles"][style.name] = style_info
        
        # Extract paragraph information
        for i, para in enumerate(doc.paragraphs):
            para_info = {
                "index": i,
                "text": para.text,
                "style": para.style.name if para.style else "Default",
                "alignment": str(para.alignment) if para.alignment else None,
                "runs": []
            }
            
            # Extract run information (formatting within paragraph)
            for j, run in enumerate(para.runs):
                run_info = {
                    "index": j,
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "font": run.font.name if run.font.name else None,
                    "font_size": run.font.size.pt if run.font.size else None
                }
                para_info["runs"].append(run_info)
            
            doc_info["paragraphs"].append(para_info)
            
            # Extract citations
            citation_pattern_numbered = r'\[\d+(?:[-–,]\d+)*\]'
            citation_pattern_author_date = r'\([A-Za-z]+(?: et al\.)?(?:, \d{4}(?:[a-z])?)+\)'
            
            numbered_citations = re.findall(citation_pattern_numbered, para.text)
            author_date_citations = re.findall(citation_pattern_author_date, para.text)
            
            if numbered_citations:
                for citation in numbered_citations:
                    doc_info["citations"].append({
                        "text": citation,
                        "type": "numbered",
                        "paragraph_index": i
                    })
            
            if author_date_citations:
                for citation in author_date_citations:
                    doc_info["citations"].append({
                        "text": citation,
                        "type": "author-date",
                        "paragraph_index": i
                    })
            
            # Check if paragraph is a reference
            if "references" in para.text.lower() or "bibliography" in para.text.lower():
                if len(para.text.split()) < 5:  # Likely a heading
                    doc_info["references"].append({
                        "type": "heading",
                        "text": para.text,
                        "paragraph_index": i
                    })
            
            # Check for reference entries
            if i > 0 and doc_info["paragraphs"][i-1]["text"].lower().startswith("reference"):
                if re.match(r'^\[\d+\]', para.text) or re.match(r'^\d+\.', para.text):
                    doc_info["references"].append({
                        "type": "entry",
                        "text": para.text,
                        "paragraph_index": i,
                        "format": "numbered" if re.match(r'^\[\d+\]', para.text) else "decimal"
                    })
            
            # Extract author information (usually in the first few paragraphs)
            if i < 5:
                # Look for author names and affiliations
                if re.search(r'^[A-Z][a-z]+ [A-Z][a-z]+', para.text) and len(para.text.split()) < 10:
                    doc_info["authors"].append({
                        "text": para.text,
                        "paragraph_index": i
                    })
                # Look for affiliations
                if re.search(r'Department|University|Institute|College', para.text) and len(para.text.split()) < 20:
                    doc_info["affiliations"].append({
                        "text": para.text,
                        "paragraph_index": i
                    })
        
        # Extract section information
        for i, section in enumerate(doc.sections):
            section_info = {
                "index": i,
                "page_width": section.page_width.inches,
                "page_height": section.page_height.inches,
                "left_margin": section.left_margin.inches,
                "right_margin": section.right_margin.inches,
                "top_margin": section.top_margin.inches,
                "bottom_margin": section.bottom_margin.inches,
                "header_distance": section.header_distance.inches,
                "footer_distance": section.footer_distance.inches,
                "orientation": "portrait" if section.orientation == 0 else "landscape"
            }
            doc_info["sections"].append(section_info)
        
        # Extract table information
        for i, table in enumerate(doc.tables):
            table_info = {
                "index": i,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "cells": [],
                "caption": None,
                "caption_position": None
            }
            
            # Look for table caption (usually before or after the table)
            if i > 0 and "Table" in doc.paragraphs[i-1].text and len(doc.paragraphs[i-1].text.split()) < 20:
                table_info["caption"] = doc.paragraphs[i-1].text
                table_info["caption_position"] = "before"
            elif i < len(doc.paragraphs) - 1 and "Table" in doc.paragraphs[i+1].text and len(doc.paragraphs[i+1].text.split()) < 20:
                table_info["caption"] = doc.paragraphs[i+1].text
                table_info["caption_position"] = "after"
            
            for r, row in enumerate(table.rows):
                for c, cell in enumerate(row.cells):
                    cell_info = {
                        "row": r,
                        "column": c,
                        "text": cell.text,
                        "paragraphs": []
                    }
                    
                    for p, para in enumerate(cell.paragraphs):
                        para_info = {
                            "index": p,
                            "text": para.text,
                            "style": para.style.name if para.style else "Default"
                        }
                        cell_info["paragraphs"].append(para_info)
                    
                    table_info["cells"].append(cell_info)
            
            doc_info["tables"].append(table_info)
        
        # Look for figures (approximation based on text patterns)
        for i, para in enumerate(doc.paragraphs):
            if "Figure" in para.text and len(para.text.split()) < 20:
                figure_info = {
                    "index": i,
                    "caption": para.text,
                    "caption_position": "unknown"  # Hard to determine without image analysis
                }
                doc_info["figures"].append(figure_info)
        
        return doc_info
    
    def analyze_documents(self):
        """Extract information from both template and article documents."""
        self.template_info = self.extract_document_info(self.template_file, "template")
        self.article_info = self.extract_document_info(self.article_file, "article")
    
    def extract_sections(self, doc_info):
        """Extract main sections from document."""
        sections = []
        
        # Find title
        if doc_info["paragraphs"]:
            title = doc_info["paragraphs"][0]["text"]
            sections.append(f"Title: \"{title}\"")
        
        # Look for common sections in academic papers
        section_keywords = ["ABSTRACT", "Introduction", "Related Work", "Method", "Materials and Methods",
                           "Results", "Discussion", "Conclusion", "References"]
        
        for keyword in section_keywords:
            for para in doc_info["paragraphs"]:
                if keyword.lower() in para["text"].lower() and len(para["text"]) < 50:
                    sections.append(f"{keyword}: Present")
                    break
        
        return sections
    
    def extract_formatting_requirements(self, doc_info):
        """Extract formatting requirements from template."""
        formatting = []
        
        # Font information
        fonts = set()
        for style_name, style_info in doc_info["styles"].items():
            if style_info["font"]:
                fonts.add(style_info["font"])
        
        if fonts:
            formatting.append(f"Font: Primary fonts are {', '.join(fonts)}")
        
        # Paragraph alignment
        alignments = {}
        for para in doc_info["paragraphs"]:
            if para["alignment"]:
                alignments[para["alignment"]] = alignments.get(para["alignment"], 0) + 1
        
        if alignments:
            main_alignment = max(alignments.items(), key=lambda x: x[1])[0]
            formatting.append(f"Paragraph Alignment: {main_alignment}")
        
        # Heading styles
        heading_styles = {}
        for style_name, style_info in doc_info["styles"].items():
            if "heading" in style_name.lower():
                heading_styles[style_name] = style_info
        
        if heading_styles:
            formatting.append("Headings: Various levels with specific formatting")
        
        # Tables
        if doc_info["tables_count"] > 0:
            formatting.append(f"Tables: {doc_info['tables_count']} tables with specific formatting")
        
        # Citations
        citation_types = {"numbered": 0, "author-date": 0}
        for citation in doc_info["citations"]:
            citation_types[citation["type"]] += 1
        
        if citation_types["numbered"] > citation_types["author-date"]:
            formatting.append("Citations: Numbered format [1], [2, 3], [4-6]")
        elif citation_types["author-date"] > citation_types["numbered"]:
            formatting.append("Citations: Author-date format (Smith et al., 2020)")
        
        return formatting
    
    def identify_potential_issues(self):
        """Identify potential issues in the article format."""
        issues = []
        
        # Check for grammatical and typographical errors
        error_count = 0
        for para in self.article_info["paragraphs"]:
            # Simple check for common errors
            text = para["text"]
            if "  " in text:  # Double spaces
                error_count += 1
            if re.search(r'[a-z][A-Z]', text):  # Missing space between sentences
                error_count += 1
            if re.search(r'[.,;:]\w', text):  # Missing space after punctuation
                error_count += 1
        
        if error_count > 10:
            issues.append(f"Potential grammatical and typographical errors ({error_count} instances detected)")
        
        # Check for inconsistent formatting
        style_counts = {}
        for para in self.article_info["paragraphs"]:
            style_counts[para["style"]] = style_counts.get(para["style"], 0) + 1
        
        if len(style_counts) > 10:
            issues.append(f"Inconsistent paragraph styles ({len(style_counts)} different styles detected)")
        
        # Check for citation format consistency
        citation_types = {"numbered": 0, "author-date": 0}
        for citation in self.article_info["citations"]:
            citation_types[citation["type"]] += 1
        
        if citation_types["numbered"] > 0 and citation_types["author-date"] > 0:
            issues.append(f"Inconsistent citation format (both numbered [{citation_types['numbered']}] and author-date [{citation_types['author-date']}] detected)")
        
        # Check for table and figure caption consistency
        table_captions = [table.get("caption") for table in self.article_info["tables"] if table.get("caption")]
        if len(table_captions) < self.article_info["tables_count"]:
            issues.append(f"Missing table captions ({self.article_info['tables_count'] - len(table_captions)} tables without captions)")
        
        # Check for section numbering consistency
        section_numbering_patterns = set()
        for para in self.article_info["paragraphs"]:
            if para["style"].lower().startswith("heading"):
                # Check for different numbering patterns
                if re.match(r'^[IVX]+\.', para["text"]):  # Roman numerals
                    section_numbering_patterns.add("roman")
                elif re.match(r'^\d+\.', para["text"]):  # Decimal
                    section_numbering_patterns.add("decimal")
                elif re.match(r'^\d+\.\d+', para["text"]):  # Multi-level decimal
                    section_numbering_patterns.add("multi-decimal")
        
        if len(section_numbering_patterns) > 1:
            issues.append(f"Inconsistent section numbering patterns ({', '.join(section_numbering_patterns)})")
        
        return issues
    
    def compare_documents(self):
        """Compare template and article to identify differences."""
        # Compare document structure
        self.compare_structure()
        
        # Compare formatting
        self.compare_formatting()
        
        # Compare content organization
        self.compare_content_organization()
        
        # Compare citation styles
        self.compare_citation_styles()
        
        # Compare table and figure formatting
        self.compare_table_figure_formatting()
        
        # Compare reference formatting
        self.compare_reference_formatting()
        
        # Compare section numbering
        self.compare_section_numbering()
        
        # Analyze academic quality
        self.analyze_academic_quality()
        
        # Check for Qubahan Academic Journal citation
        self.check_journal_citation()
        
        # Check for self-citation
        self.check_self_citation()
        
        # Identify major issues
        self.identify_major_issues()
    
    def compare_structure(self):
        """Compare document structure between template and article."""
        # Extract sections from both documents
        template_sections = [para["text"] for para in self.template_info["paragraphs"] 
                            if para["style"].lower().startswith("heading") or para["text"].isupper()]
        
        article_sections = [para["text"] for para in self.article_info["paragraphs"] 
                           if para["style"].lower().startswith("heading") or para["text"].isupper()]
        
        # Find common and different sections
        matcher = difflib.SequenceMatcher(None, template_sections, article_sections)
        
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                for section in template_sections[i1:i2]:
                    self.comparison_results["structure"].append({
                        "element": section,
                        "status": "compliant",
                        "details": "Section present in both documents"
                    })
            elif tag == 'delete':
                for section in template_sections[i1:i2]:
                    self.comparison_results["structure"].append({
                        "element": section,
                        "status": "missing",
                        "details": "Required section missing in article"
                    })
            elif tag == 'insert':
                for section in article_sections[j1:j2]:
                    self.comparison_results["structure"].append({
                        "element": section,
                        "status": "extra",
                        "details": "Extra section in article not in template"
                    })
            elif tag == 'replace':
                for section in template_sections[i1:i2]:
                    self.comparison_results["structure"].append({
                        "element": section,
                        "status": "missing",
                        "details": "Required section missing in article"
                    })
                for section in article_sections[j1:j2]:
                    self.comparison_results["structure"].append({
                        "element": section,
                        "status": "extra",
                        "details": "Extra section in article not in template"
                    })
    
    def compare_formatting(self):
        """Compare formatting between template and article."""
        # Compare fonts
        template_fonts = set()
        article_fonts = set()
        
        for style_name, style_info in self.template_info["styles"].items():
            if style_info["font"]:
                template_fonts.add(style_info["font"])
        
        for style_name, style_info in self.article_info["styles"].items():
            if style_info["font"]:
                article_fonts.add(style_info["font"])
        
        if template_fonts == article_fonts:
            self.comparison_results["formatting"].append({
                "element": "Fonts",
                "status": "compliant",
                "details": f"Both documents use the same fonts: {', '.join(template_fonts)}"
            })
        else:
            self.comparison_results["formatting"].append({
                "element": "Fonts",
                "status": "non-compliant",
                "details": f"Template fonts: {', '.join(template_fonts)}; Article fonts: {', '.join(article_fonts)}"
            })
        
        # Compare paragraph alignment
        template_alignments = {}
        article_alignments = {}
        
        for para in self.template_info["paragraphs"]:
            if para["alignment"]:
                template_alignments[para["alignment"]] = template_alignments.get(para["alignment"], 0) + 1
        
        for para in self.article_info["paragraphs"]:
            if para["alignment"]:
                article_alignments[para["alignment"]] = article_alignments.get(para["alignment"], 0) + 1
        
        if template_alignments and article_alignments:
            template_main = max(template_alignments.items(), key=lambda x: x[1])[0]
            article_main = max(article_alignments.items(), key=lambda x: x[1])[0]
            
            if template_main == article_main:
                self.comparison_results["formatting"].append({
                    "element": "Paragraph Alignment",
                    "status": "compliant",
                    "details": f"Both documents primarily use {template_main} alignment"
                })
            else:
                self.comparison_results["formatting"].append({
                    "element": "Paragraph Alignment",
                    "status": "non-compliant",
                    "details": f"Template primarily uses {template_main}; Article primarily uses {article_main}"
                })
        
        # Compare tables
        if self.template_info["tables_count"] == self.article_info["tables_count"]:
            self.comparison_results["formatting"].append({
                "element": "Tables Count",
                "status": "compliant",
                "details": f"Both documents have {self.template_info['tables_count']} tables"
            })
        else:
            self.comparison_results["formatting"].append({
                "element": "Tables Count",
                "status": "non-compliant",
                "details": f"Template has {self.template_info['tables_count']} tables; Article has {self.article_info['tables_count']} tables"
            })
    
    def compare_content_organization(self):
        """Compare content organization between template and article."""
        # Check for abstract
        template_abstract = None
        article_abstract = None
        
        for para in self.template_info["paragraphs"]:
            if para["text"].startswith("ABSTRACT:") or para["text"].startswith("Abstract:"):
                template_abstract = para["text"]
                break
        
        for para in self.article_info["paragraphs"]:
            if para["text"].startswith("ABSTRACT") or para["text"].startswith("Abstract"):
                article_abstract = para["text"]
                break
        
        if template_abstract and article_abstract:
            template_words = len(template_abstract.split())
            article_words = len(article_abstract.split())
            
            if 150 <= article_words <= 300:
                self.comparison_results["content"].append({
                    "element": "Abstract Length",
                    "status": "compliant",
                    "details": f"Article abstract has {article_words} words (within 150-300 word limit)"
                })
            else:
                self.comparison_results["content"].append({
                    "element": "Abstract Length",
                    "status": "non-compliant",
                    "details": f"Article abstract has {article_words} words (outside 150-300 word limit)"
                })
        
        # Check for keywords
        template_keywords = None
        article_keywords = None
        
        for para in self.template_info["paragraphs"]:
            if para["text"].startswith("Keywords:"):
                template_keywords = para["text"]
                break
        
        for para in self.article_info["paragraphs"]:
            if para["text"].startswith("Keywords:"):
                article_keywords = para["text"]
                break
        
        if template_keywords and article_keywords:
            template_count = len(template_keywords.split(","))
            article_count = len(article_keywords.split(","))
            
            if template_count == article_count:
                self.comparison_results["content"].append({
                    "element": "Keywords Count",
                    "status": "compliant",
                    "details": f"Both documents have {template_count} keywords"
                })
            else:
                self.comparison_results["content"].append({
                    "element": "Keywords Count",
                    "status": "non-compliant",
                    "details": f"Template has {template_count} keywords; Article has {article_count} keywords"
                })
    
    def compare_citation_styles(self):
        """Compare citation styles between template and article."""
        # Count citation types in template
        template_citation_types = {"numbered": 0, "author-date": 0}
        for citation in self.template_info["citations"]:
            template_citation_types[citation["type"]] += 1
        
        # Count citation types in article
        article_citation_types = {"numbered": 0, "author-date": 0}
        for citation in self.article_info["citations"]:
            article_citation_types[citation["type"]] += 1
        
        # Determine primary citation style in template
        template_style = "numbered" if template_citation_types["numbered"] > template_citation_types["author-date"] else "author-date"
        
        # Determine primary citation style in article
        article_style = "numbered" if article_citation_types["numbered"] > article_citation_types["author-date"] else "author-date"
        
        # Compare citation styles
        if template_style == article_style:
            self.comparison_results["citations"].append({
                "element": "Citation Style",
                "status": "compliant",
                "details": f"Both documents use {template_style} citation style"
            })
        else:
            self.comparison_results["citations"].append({
                "element": "Citation Style",
                "status": "non-compliant",
                "details": f"Template uses {template_style} citation style; Article uses {article_style} citation style"
            })
        
        # Check for citation consistency in article
        if article_citation_types["numbered"] > 0 and article_citation_types["author-date"] > 0:
            self.comparison_results["citations"].append({
                "element": "Citation Consistency",
                "status": "non-compliant",
                "details": f"Article uses mixed citation styles: {article_citation_types['numbered']} numbered and {article_citation_types['author-date']} author-date citations"
            })
        else:
            self.comparison_results["citations"].append({
                "element": "Citation Consistency",
                "status": "compliant",
                "details": f"Article consistently uses {article_style} citation style"
            })
        
        # Check for citation format in numbered style
        if article_style == "numbered":
            # Check for proper formatting of numbered citations
            proper_format = True
            improper_examples = []
            
            for citation in self.article_info["citations"]:
                if citation["type"] == "numbered":
                    # Check if citation follows [n] format
                    if not re.match(r'^\[\d+(?:[-–,]\d+)*\]$', citation["text"]):
                        proper_format = False
                        improper_examples.append(citation["text"])
            
            if proper_format:
                self.comparison_results["citations"].append({
                    "element": "Numbered Citation Format",
                    "status": "compliant",
                    "details": "All numbered citations follow proper format [n]"
                })
            else:
                self.comparison_results["citations"].append({
                    "element": "Numbered Citation Format",
                    "status": "non-compliant",
                    "details": f"Some numbered citations do not follow proper format, examples: {', '.join(improper_examples[:3])}"
                })
    
    def compare_table_figure_formatting(self):
        """Compare table and figure formatting between template and article."""
        # Check table captions in template
        template_tables_with_captions = sum(1 for table in self.template_info["tables"] if table.get("caption"))
        
        # Check table captions in article
        article_tables_with_captions = sum(1 for table in self.article_info["tables"] if table.get("caption"))
        
        # Compare table caption presence
        if self.article_info["tables_count"] > 0:
            if article_tables_with_captions == self.article_info["tables_count"]:
                self.comparison_results["tables_figures"].append({
                    "element": "Table Captions",
                    "status": "compliant",
                    "details": "All tables have captions"
                })
            else:
                self.comparison_results["tables_figures"].append({
                    "element": "Table Captions",
                    "status": "non-compliant",
                    "details": f"{self.article_info['tables_count'] - article_tables_with_captions} out of {self.article_info['tables_count']} tables missing captions"
                })
        
        # Check table caption format
        if article_tables_with_captions > 0:
            # Check if captions follow "Table X: Description" format
            proper_format = True
            improper_examples = []
            
            for table in self.article_info["tables"]:
                if table.get("caption"):
                    if not re.match(r'^Table \d+[.:] .+', table["caption"]):
                        proper_format = False
                        improper_examples.append(table["caption"])
            
            if proper_format:
                self.comparison_results["tables_figures"].append({
                    "element": "Table Caption Format",
                    "status": "compliant",
                    "details": "All table captions follow proper format 'Table X: Description'"
                })
            else:
                self.comparison_results["tables_figures"].append({
                    "element": "Table Caption Format",
                    "status": "non-compliant",
                    "details": f"Some table captions do not follow proper format, examples: {', '.join(improper_examples[:2])}"
                })
        
        # Check figure captions
        if self.article_info["figures"]:
            # Check if figure captions follow "Figure X: Description" format
            proper_format = True
            improper_examples = []
            
            for figure in self.article_info["figures"]:
                if not re.match(r'^Figure \d+[.:] .+', figure["caption"]):
                    proper_format = False
                    improper_examples.append(figure["caption"])
            
            if proper_format:
                self.comparison_results["tables_figures"].append({
                    "element": "Figure Caption Format",
                    "status": "compliant",
                    "details": "All figure captions follow proper format 'Figure X: Description'"
                })
            else:
                self.comparison_results["tables_figures"].append({
                    "element": "Figure Caption Format",
                    "status": "non-compliant",
                    "details": f"Some figure captions do not follow proper format, examples: {', '.join(improper_examples[:2])}"
                })
    
    def compare_reference_formatting(self):
        """Compare reference formatting between template and article."""
        # Check if references section exists
        template_has_references = any(ref["type"] == "heading" for ref in self.template_info["references"])
        article_has_references = any(ref["type"] == "heading" for ref in self.article_info["references"])
        
        if template_has_references and article_has_references:
            self.comparison_results["references"].append({
                "element": "References Section",
                "status": "compliant",
                "details": "References section present in both documents"
            })
        elif not article_has_references:
            self.comparison_results["references"].append({
                "element": "References Section",
                "status": "non-compliant",
                "details": "References section missing in article"
            })
        
        # Check reference entry format
        template_reference_format = None
        article_reference_format = None
        
        for ref in self.template_info["references"]:
            if ref["type"] == "entry" and "format" in ref:
                template_reference_format = ref["format"]
                break
        
        for ref in self.article_info["references"]:
            if ref["type"] == "entry" and "format" in ref:
                article_reference_format = ref["format"]
                break
        
        if template_reference_format and article_reference_format:
            if template_reference_format == article_reference_format:
                self.comparison_results["references"].append({
                    "element": "Reference Format",
                    "status": "compliant",
                    "details": f"Both documents use {template_reference_format} reference format"
                })
            else:
                self.comparison_results["references"].append({
                    "element": "Reference Format",
                    "status": "non-compliant",
                    "details": f"Template uses {template_reference_format} reference format; Article uses {article_reference_format} reference format"
                })
    
    def compare_section_numbering(self):
        """Compare section numbering between template and article."""
        # Extract section numbering patterns from template
        template_patterns = set()
        for para in self.template_info["paragraphs"]:
            if para["style"].lower().startswith("heading"):
                if re.match(r'^[IVX]+\.', para["text"]):  # Roman numerals
                    template_patterns.add("roman")
                elif re.match(r'^\d+\.', para["text"]):  # Decimal
                    template_patterns.add("decimal")
                elif re.match(r'^\d+\.\d+', para["text"]):  # Multi-level decimal
                    template_patterns.add("multi-decimal")
        
        # Extract section numbering patterns from article
        article_patterns = set()
        for para in self.article_info["paragraphs"]:
            if para["style"].lower().startswith("heading"):
                if re.match(r'^[IVX]+\.', para["text"]):  # Roman numerals
                    article_patterns.add("roman")
                elif re.match(r'^\d+\.', para["text"]):  # Decimal
                    article_patterns.add("decimal")
                elif re.match(r'^\d+\.\d+', para["text"]):  # Multi-level decimal
                    article_patterns.add("multi-decimal")
        
        # Compare section numbering patterns
        if template_patterns and article_patterns:
            if template_patterns == article_patterns:
                self.comparison_results["section_numbering"].append({
                    "element": "Section Numbering Pattern",
                    "status": "compliant",
                    "details": f"Both documents use {', '.join(template_patterns)} section numbering"
                })
            else:
                self.comparison_results["section_numbering"].append({
                    "element": "Section Numbering Pattern",
                    "status": "non-compliant",
                    "details": f"Template uses {', '.join(template_patterns)} numbering; Article uses {', '.join(article_patterns)} numbering"
                })
        
        # Check for section numbering consistency in article
        if len(article_patterns) > 1:
            self.comparison_results["section_numbering"].append({
                "element": "Section Numbering Consistency",
                "status": "non-compliant",
                "details": f"Article uses inconsistent section numbering patterns: {', '.join(article_patterns)}"
            })
        elif len(article_patterns) == 1:
            self.comparison_results["section_numbering"].append({
                "element": "Section Numbering Consistency",
                "status": "compliant",
                "details": f"Article consistently uses {next(iter(article_patterns))} section numbering"
            })
    
    def analyze_academic_quality(self):
        """Analyze academic quality metrics."""
        # Calculate sentence length variation
        sentences = []
        for para in self.article_info["paragraphs"]:
            # Simple sentence splitting
            para_sentences = re.split(r'[.!?]+', para["text"])
            sentences.extend([s.strip() for s in para_sentences if s.strip()])
        
        sentence_lengths = [len(s.split()) for s in sentences]
        
        if sentence_lengths:
            avg_length = sum(sentence_lengths) / len(sentence_lengths)
            min_length = min(sentence_lengths)
            max_length = max(sentence_lengths)
            
            if min_length < 5 and max_length > 40:
                self.comparison_results["academic_quality"].append({
                    "element": "Sentence Length Variation",
                    "status": "non-compliant",
                    "details": f"High sentence length variation (min: {min_length}, max: {max_length}, avg: {avg_length:.1f})"
                })
            elif 10 <= avg_length <= 25:
                self.comparison_results["academic_quality"].append({
                    "element": "Sentence Length Variation",
                    "status": "compliant",
                    "details": f"Appropriate sentence length variation (min: {min_length}, max: {max_length}, avg: {avg_length:.1f})"
                })
            else:
                self.comparison_results["academic_quality"].append({
                    "element": "Sentence Length Variation",
                    "status": "partially-compliant",
                    "details": f"Suboptimal sentence length variation (min: {min_length}, max: {max_length}, avg: {avg_length:.1f})"
                })
        
        # Check for passive voice usage
        passive_indicators = ["is", "are", "was", "were", "be", "been", "being"]
        passive_count = 0
        
        for sentence in sentences:
            words = sentence.lower().split()
            if any(indicator in words for indicator in passive_indicators) and "by" in words:
                passive_count += 1
        
        passive_percentage = (passive_count / len(sentences) * 100) if sentences else 0
        
        if passive_percentage > 30:
            self.comparison_results["academic_quality"].append({
                "element": "Passive Voice Usage",
                "status": "non-compliant",
                "details": f"Excessive passive voice usage ({passive_percentage:.1f}% of sentences)"
            })
        elif passive_percentage > 15:
            self.comparison_results["academic_quality"].append({
                "element": "Passive Voice Usage",
                "status": "partially-compliant",
                "details": f"Moderate passive voice usage ({passive_percentage:.1f}% of sentences)"
            })
        else:
            self.comparison_results["academic_quality"].append({
                "element": "Passive Voice Usage",
                "status": "compliant",
                "details": f"Appropriate passive voice usage ({passive_percentage:.1f}% of sentences)"
            })
        
        # Check for paragraph length consistency
        paragraph_lengths = [len(para["text"].split()) for para in self.article_info["paragraphs"] if para["text"].strip()]
        
        if paragraph_lengths:
            avg_para_length = sum(paragraph_lengths) / len(paragraph_lengths)
            min_para_length = min(paragraph_lengths)
            max_para_length = max(paragraph_lengths)
            
            if min_para_length < 20 and max_para_length > 300:
                self.comparison_results["academic_quality"].append({
                    "element": "Paragraph Length Consistency",
                    "status": "non-compliant",
                    "details": f"Inconsistent paragraph lengths (min: {min_para_length}, max: {max_para_length}, avg: {avg_para_length:.1f})"
                })
            elif 50 <= avg_para_length <= 200:
                self.comparison_results["academic_quality"].append({
                    "element": "Paragraph Length Consistency",
                    "status": "compliant",
                    "details": f"Appropriate paragraph lengths (min: {min_para_length}, max: {max_para_length}, avg: {avg_para_length:.1f})"
                })
            else:
                self.comparison_results["academic_quality"].append({
                    "element": "Paragraph Length Consistency",
                    "status": "partially-compliant",
                    "details": f"Suboptimal paragraph lengths (min: {min_para_length}, max: {max_para_length}, avg: {avg_para_length:.1f})"
                })
    
    def check_journal_citation(self):
        """Check if Qubahan Academic Journal is cited in the document."""
        # Look for Qubahan Academic Journal citation in text
        qaj_cited_in_text = False
        qaj_cited_in_references = False
        
        # Check in text
        for para in self.article_info["paragraphs"]:
            if re.search(r'Qubahan Academic Journal|QAJ', para["text"], re.IGNORECASE):
                qaj_cited_in_text = True
                break
        
        # Check in references
        for ref in self.article_info["references"]:
            if ref["type"] == "entry" and re.search(r'Qubahan Academic Journal|QAJ', ref["text"], re.IGNORECASE):
                qaj_cited_in_references = True
                break
        
        # Add results to comparison
        if qaj_cited_in_text:
            self.comparison_results["journal_citation"].append({
                "element": "QAJ Citation in Text",
                "status": "compliant",
                "details": "Qubahan Academic Journal is cited in the text"
            })
        else:
            self.comparison_results["journal_citation"].append({
                "element": "QAJ Citation in Text",
                "status": "non-compliant",
                "details": "Qubahan Academic Journal is not cited in the text"
            })
        
        if qaj_cited_in_references:
            self.comparison_results["journal_citation"].append({
                "element": "QAJ Citation in References",
                "status": "compliant",
                "details": "Qubahan Academic Journal is cited in the references"
            })
        else:
            self.comparison_results["journal_citation"].append({
                "element": "QAJ Citation in References",
                "status": "non-compliant",
                "details": "Qubahan Academic Journal is not cited in the references"
            })
    
    def check_self_citation(self):
        """Check for self-citation in the document."""
        # Extract author names from the document
        author_names = []
        for author in self.article_info["authors"]:
            # Extract last names
            name_parts = author["text"].split()
            if name_parts:
                author_names.append(name_parts[-1])  # Last name
        
        # Check for self-citation in references
        self_citations = []
        
        for ref in self.article_info["references"]:
            if ref["type"] == "entry":
                for author_name in author_names:
                    if author_name in ref["text"]:
                        self_citations.append(ref["text"])
        
        # Add results to comparison
        if self_citations:
            self.comparison_results["self_citation"].append({
                "element": "Self-Citation",
                "status": "non-compliant",
                "details": f"Self-citation detected: {len(self_citations)} instances found"
            })
            
            # Add examples of self-citations
            if len(self_citations) > 0:
                examples = self_citations[:2]  # Limit to 2 examples
                self.comparison_results["self_citation"].append({
                    "element": "Self-Citation Examples",
                    "status": "non-compliant",
                    "details": f"Examples: {'; '.join(examples)}"
                })
        else:
            self.comparison_results["self_citation"].append({
                "element": "Self-Citation",
                "status": "compliant",
                "details": "No self-citation detected in the document"
            })
    
    def identify_major_issues(self):
        """Identify major issues based on comparison results."""
        # Structure issues
        missing_sections = [item["element"] for item in self.comparison_results["structure"] 
                           if item["status"] == "missing"]
        
        if missing_sections:
            self.comparison_results["major_issues"].append({
                "category": "Structure",
                "issue": f"Missing required sections: {', '.join(missing_sections[:3])}{'...' if len(missing_sections) > 3 else ''}"
            })
        
        # Formatting issues
        non_compliant_formatting = [item["element"] for item in self.comparison_results["formatting"] 
                                   if item["status"] == "non-compliant"]
        
        if non_compliant_formatting:
            self.comparison_results["major_issues"].append({
                "category": "Formatting",
                "issue": f"Non-compliant formatting: {', '.join(non_compliant_formatting)}"
            })
        
        # Content issues
        non_compliant_content = [item["element"] for item in self.comparison_results["content"] 
                                if item["status"] == "non-compliant"]
        
        if non_compliant_content:
            self.comparison_results["major_issues"].append({
                "category": "Content",
                "issue": f"Non-compliant content: {', '.join(non_compliant_content)}"
            })
        
        # Citation issues
        non_compliant_citations = [item["element"] for item in self.comparison_results["citations"] 
                                  if item["status"] == "non-compliant"]
        
        if non_compliant_citations:
            self.comparison_results["major_issues"].append({
                "category": "Citations",
                "issue": f"Non-compliant citations: {', '.join(non_compliant_citations)}"
            })
        
        # Table and figure issues
        non_compliant_tables_figures = [item["element"] for item in self.comparison_results["tables_figures"] 
                                       if item["status"] == "non-compliant"]
        
        if non_compliant_tables_figures:
            self.comparison_results["major_issues"].append({
                "category": "Tables and Figures",
                "issue": f"Non-compliant tables/figures: {', '.join(non_compliant_tables_figures)}"
            })
        
        # Reference issues
        non_compliant_references = [item["element"] for item in self.comparison_results["references"] 
                                   if item["status"] == "non-compliant"]
        
        if non_compliant_references:
            self.comparison_results["major_issues"].append({
                "category": "References",
                "issue": f"Non-compliant references: {', '.join(non_compliant_references)}"
            })
        
        # Section numbering issues
        non_compliant_section_numbering = [item["element"] for item in self.comparison_results["section_numbering"] 
                                          if item["status"] == "non-compliant"]
        
        if non_compliant_section_numbering:
            self.comparison_results["major_issues"].append({
                "category": "Section Numbering",
                "issue": f"Non-compliant section numbering: {', '.join(non_compliant_section_numbering)}"
            })
        
        # Academic quality issues
        non_compliant_academic_quality = [item["element"] for item in self.comparison_results["academic_quality"] 
                                         if item["status"] == "non-compliant"]
        
        if non_compliant_academic_quality:
            self.comparison_results["major_issues"].append({
                "category": "Academic Quality",
                "issue": f"Academic quality issues: {', '.join(non_compliant_academic_quality)}"
            })
        
        # Journal citation issues
        non_compliant_journal_citation = [item["element"] for item in self.comparison_results["journal_citation"] 
                                         if item["status"] == "non-compliant"]
        
        if non_compliant_journal_citation:
            self.comparison_results["major_issues"].append({
                "category": "Journal Citation",
                "issue": f"Journal citation issues: {', '.join(non_compliant_journal_citation)}"
            })
        
        # Self-citation issues
        non_compliant_self_citation = [item["element"] for item in self.comparison_results["self_citation"] 
                                      if item["status"] == "non-compliant" and item["element"] == "Self-Citation"]
        
        if non_compliant_self_citation:
            self.comparison_results["major_issues"].append({
                "category": "Self-Citation",
                "issue": "Self-citation detected in the document"
            })
        
        # Check for grammatical issues
        error_count = 0
        for para in self.article_info["paragraphs"]:
            text = para["text"]
            if "  " in text:  # Double spaces
                error_count += 1
            if re.search(r'[a-z][A-Z]', text):  # Missing space between sentences
                error_count += 1
            if re.search(r'[.,;:]\w', text):  # Missing space after punctuation
                error_count += 1
        
        if error_count > 10:
            self.comparison_results["major_issues"].append({
                "category": "Writing Quality",
                "issue": f"Potential grammatical and typographical errors ({error_count} instances detected)"
            })
    
    def generate_compliance_chart(self):
        """Generate a chart visualizing compliance status."""
        # Count compliance status for each category
        categories = ["Structure", "Formatting", "Content", "Citations", "Tables/Figures", 
                     "References", "Section Numbering", "Academic Quality", "Journal Citation", "Self-Citation"]
        category_keys = ["structure", "formatting", "content", "citations", "tables_figures", 
                        "references", "section_numbering", "academic_quality", "journal_citation", "self_citation"]
        
        compliant = []
        non_compliant = []
        partially_compliant = []
        
        for category in category_keys:
            category_items = self.comparison_results[category]
            if category_items:
                category_compliant = sum(1 for item in category_items if item.get("status") == "compliant")
                category_non_compliant = sum(1 for item in category_items if item.get("status") == "non-compliant")
                category_partially = sum(1 for item in category_items if item.get("status") == "partially-compliant")
                category_total = len(category_items)
                
                compliant.append(category_compliant / category_total * 100)
                non_compliant.append(category_non_compliant / category_total * 100)
                partially_compliant.append(category_partially / category_total * 100)
            else:
                compliant.append(0)
                non_compliant.append(0)
                partially_compliant.append(0)
        
        # Create chart
        fig, ax = plt.subplots(figsize=(14, 7))
        
        x = np.arange(len(categories))
        width = 0.25
        
        ax.bar(x - width, compliant, width, label='Compliant', color='#4CAF50')
        ax.bar(x, partially_compliant, width, label='Partially Compliant', color='#FFC107')
        ax.bar(x + width, non_compliant, width, label='Non-Compliant', color='#F44336')
        
        ax.set_title('Compliance by Category')
        ax.set_ylabel('Percentage')
        ax.set_yticks(np.arange(0, 101, 20))
        ax.set_xticks(x)
        ax.set_xticklabels(categories, rotation=45, ha='right')
        ax.legend()
        
        plt.tight_layout()
        
        # Convert plot to image
        buf = BytesIO()
        plt.savefig(buf, format='png')
        plt.close(fig)
        buf.seek(0)
        
        return buf
    
    def run_analysis(self):
        """Run the complete document analysis and comparison."""
        self.analyze_documents()
        self.compare_documents()
        return self.comparison_results

def main():
    st.title("QAJ Format Checker by Ridwan Marqas")
    
    st.markdown("""
    This tool compares documents against the QAJ template to check for format compliance. 
    Upload a QAJ template document and an article document to analyze their structure, formatting, and content.
    """)
    
    # File upload
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Template Document")
        template_file = st.file_uploader("Upload QAJ template document", type=["docx"], key="template")
    
    with col2:
        st.subheader("Article Document")
        article_file = st.file_uploader("Upload article document to check", type=["docx"], key="article")
    
    if template_file and article_file:
        # Run analysis
        with st.spinner("Analyzing documents..."):
            analyzer = DocumentAnalyzer(template_file, article_file)
            comparison_results = analyzer.run_analysis()
            
            # Generate compliance chart
            chart_image = analyzer.generate_compliance_chart()
        
        # Display results
        st.success("Analysis complete!")
        
        # Calculate overall compliance
        total_items = 0
        compliant_items = 0
        partially_compliant_items = 0
        
        for category in ["structure", "formatting", "content", "citations", "tables_figures", 
                         "references", "section_numbering", "academic_quality", "journal_citation", "self_citation"]:
            category_items = comparison_results[category]
            total_items += len(category_items)
            compliant_items += sum(1 for item in category_items if item.get("status") == "compliant")
            partially_compliant_items += sum(1 for item in category_items if item.get("status") == "partially-compliant")
        
        compliance_percentage = (compliant_items / total_items * 100) if total_items > 0 else 0
        partial_percentage = (partially_compliant_items / total_items * 100) if total_items > 0 else 0
        
        # Display compliance score
        st.subheader("Overall Compliance")
        
        col1, col2, col3 = st.columns([1, 1, 2])
        
        with col1:
            st.metric("Compliance Score", f"{compliance_percentage:.1f}%")
        
        with col2:
            st.metric("Partial Compliance", f"{partial_percentage:.1f}%")
        
        with col3:
            if compliance_percentage >= 80:
                st.success("The document is **largely compliant** with the QAJ template requirements.")
            elif compliance_percentage >= 50:
                st.warning("The document is **partially compliant** with the QAJ template requirements.")
            else:
                st.error("The document is **mostly non-compliant** with the QAJ template requirements.")
        
        # Check for critical issues (QAJ citation and self-citation)
        critical_issues = []
        
        # Check QAJ citation
        qaj_citation_issues = [item for item in comparison_results["journal_citation"] 
                              if item["status"] == "non-compliant"]
        if qaj_citation_issues:
            critical_issues.append("Qubahan Academic Journal is not cited in the document")
        
        # Check self-citation
        self_citation_issues = [item for item in comparison_results["self_citation"] 
                               if item["status"] == "non-compliant" and item["element"] == "Self-Citation"]
        if self_citation_issues:
            critical_issues.append("Self-citation detected in the document")
        
        # Display critical issues
        if critical_issues:
            st.markdown("<div class='critical-issue'><strong>Critical Issues:</strong><ul>" + 
                       "".join([f"<li>{issue}</li>" for issue in critical_issues]) + 
                       "</ul></div>", unsafe_allow_html=True)
        
        # Display compliance chart
        st.image(chart_image, caption="Compliance by Category", use_column_width=True)
        
        # Create tabs for different sections
        tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9, tab10, tab11 = st.tabs([
            "Major Issues", "Structure", "Formatting", "Content", 
            "Citations", "Tables & Figures", "References", 
            "Section Numbering", "Academic Quality", "Journal Citation", "Self-Citation"
        ])
        
        with tab1:
            st.subheader("Major Issues")
            
            if comparison_results["major_issues"]:
                for issue in comparison_results["major_issues"]:
                    st.markdown(f"### {issue['category']}")
                    st.markdown(f"{issue['issue']}")
            else:
                st.success("No major issues detected.")
        
        with tab2:
            st.subheader("Structure Comparison")
            
            # Group by status
            structure_by_status = {"compliant": [], "missing": [], "extra": []}
            for item in comparison_results["structure"]:
                if item["status"] in structure_by_status:
                    structure_by_status[item["status"]].append(item)
            
            # Compliant sections
            if structure_by_status["compliant"]:
                st.markdown("### Compliant Sections")
                for item in structure_by_status["compliant"]:
                    st.markdown(f"- **{item['element']}**: {item['details']}")
            
            # Missing sections
            if structure_by_status["missing"]:
                st.markdown("### Missing Sections")
                st.markdown("The following required sections are missing in the article:")
                for item in structure_by_status["missing"]:
                    st.markdown(f"- **{item['element']}**: {item['details']}")
            
            # Extra sections
            if structure_by_status["extra"]:
                st.markdown("### Extra Sections")
                st.markdown("The following sections in the article are not present in the template:")
                for item in structure_by_status["extra"]:
                    st.markdown(f"- **{item['element']}**: {item['details']}")
        
        with tab3:
            st.subheader("Formatting Comparison")
            
            # Create table
            table_data = []
            for item in comparison_results["formatting"]:
                status_symbol = "✅" if item["status"] == "compliant" else "❌"
                status_class = "compliance-compliant" if item["status"] == "compliant" else "compliance-non-compliant"
                table_data.append([
                    item["element"], 
                    f'<span class="{status_class}">{status_symbol}</span>', 
                    item["details"]
                ])
            
            # Display table
            st.markdown(
                tabulate(table_data, headers=["Element", "Status", "Details"], tablefmt="html"),
                unsafe_allow_html=True
            )
        
        with tab4:
            st.subheader("Content Comparison")
            
            # Create table
            table_data = []
            for item in comparison_results["content"]:
                status_symbol = "✅" if item["status"] == "compliant" else "❌"
                status_class = "compliance-compliant" if item["status"] == "compliant" else "compliance-non-compliant"
                table_data.append([
                    item["element"], 
                    f'<span class="{status_class}">{status_symbol}</span>', 
                    item["details"]
                ])
            
            # Display table
            st.markdown(
                tabulate(table_data, headers=["Element", "Status", "Details"], tablefmt="html"),
                unsafe_allow_html=True
            )
        
        with tab5:
            st.subheader("Citation Analysis")
            
            # Create table
            table_data = []
            for item in comparison_results["citations"]:
                if item["status"] == "compliant":
                    status_symbol = "✅"
                    status_class = "compliance-compliant"
                elif item["status"] == "partially-compliant":
                    status_symbol = "⚠️"
                    status_class = "compliance-partially-compliant"
                else:
                    status_symbol = "❌"
                    status_class = "compliance-non-compliant"
                
                table_data.append([
                    item["element"], 
                    f'<span class="{status_class}">{status_symbol}</span>', 
                    item["details"]
                ])
            
            # Display table
            st.markdown(
                tabulate(table_data, headers=["Element", "Status", "Details"], tablefmt="html"),
                unsafe_allow_html=True
            )
            
            # Citation guidance
            st.markdown("### Citation Format Guidance")
            st.markdown("""
            QAJ requires **numbered citation format**. Examples of proper citation format:
            
            - Single reference: <span class="citation-example">[1]</span>
            - Multiple references: <span class="citation-example">[1, 2, 3]</span>
            - Range of references: <span class="citation-example">[1-3]</span>
            
            Citations should be numbered consecutively in the order in which they appear in the text.
            """, unsafe_allow_html=True)
        
        with tab6:
            st.subheader("Tables & Figures Analysis")
            
            # Create table
            table_data = []
            for item in comparison_results["tables_figures"]:
                if item["status"] == "compliant":
                    status_symbol = "✅"
                    status_class = "compliance-compliant"
                elif item["status"] == "partially-compliant":
                    status_symbol = "⚠️"
                    status_class = "compliance-partially-compliant"
                else:
                    status_symbol = "❌"
                    status_class = "compliance-non-compliant"
                
                table_data.append([
                    item["element"], 
                    f'<span class="{status_class}">{status_symbol}</span>', 
                    item["details"]
                ])
            
            # Display table
            st.markdown(
                tabulate(table_data, headers=["Element", "Status", "Details"], tablefmt="html"),
                unsafe_allow_html=True
            )
            
            # Table and figure guidance
            st.markdown("### Table & Figure Format Guidance")
            st.markdown("""
            **Table Format Requirements:**
            - Tables should be numbered consecutively
            - Table captions should be placed above the table
            - Format: <span class="citation-example">Table 1: Description of the table</span>
            
            **Figure Format Requirements:**
            - Figures should be numbered consecutively
            - Figure captions should be placed below the figure
            - Format: <span class="citation-example">Figure 1: Description of the figure</span>
            """, unsafe_allow_html=True)
        
        with tab7:
            st.subheader("References Analysis")
            
            # Create table
            table_data = []
            for item in comparison_results["references"]:
                if item["status"] == "compliant":
                    status_symbol = "✅"
                    status_class = "compliance-compliant"
                elif item["status"] == "partially-compliant":
                    status_symbol = "⚠️"
                    status_class = "compliance-partially-compliant"
                else:
                    status_symbol = "❌"
                    status_class = "compliance-non-compliant"
                
                table_data.append([
                    item["element"], 
                    f'<span class="{status_class}">{status_symbol}</span>', 
                    item["details"]
                ])
            
            # Display table
            st.markdown(
                tabulate(table_data, headers=["Element", "Status", "Details"], tablefmt="html"),
                unsafe_allow_html=True
            )
            
            # Reference guidance
            st.markdown("### Reference Format Guidance")
            st.markdown("""
            QAJ requires **numbered reference format** that corresponds to the citation numbers in the text.
            
            **Journal Article Format:**
            <span class="citation-example">[1] Author A, Author B, Author C. Title of article. Journal Name. Year;Volume(Issue):Page range.</span>
            
            **Book Format:**
            <span class="citation-example">[2] Author A, Author B. Title of book. Edition. City: Publisher; Year.</span>
            
            **Book Chapter Format:**
            <span class="citation-example">[3] Author A, Author B. Title of chapter. In: Editor A, Editor B, editors. Title of book. City: Publisher; Year. p. Page range.</span>
            
            **Website Format:**
            <span class="citation-example">[4] Author/Organization. Title of webpage [Internet]. Publisher; Year [cited Date]. Available from: URL</span>
            """, unsafe_allow_html=True)
        
        with tab8:
            st.subheader("Section Numbering Analysis")
            
            # Create table
            table_data = []
            for item in comparison_results["section_numbering"]:
                if item["status"] == "compliant":
                    status_symbol = "✅"
                    status_class = "compliance-compliant"
                elif item["status"] == "partially-compliant":
                    status_symbol = "⚠️"
                    status_class = "compliance-partially-compliant"
                else:
                    status_symbol = "❌"
                    status_class = "compliance-non-compliant"
                
                table_data.append([
                    item["element"], 
                    f'<span class="{status_class}">{status_symbol}</span>', 
                    item["details"]
                ])
            
            # Display table
            st.markdown(
                tabulate(table_data, headers=["Element", "Status", "Details"], tablefmt="html"),
                unsafe_allow_html=True
            )
            
            # Section numbering guidance
            st.markdown("### Section Numbering Guidance")
            st.markdown("""
            QAJ requires consistent section numbering throughout the document.
            
            **Main Sections (Level 1):**
            - Should use Roman numerals: <span class="citation-example">I. Introduction</span>
            
            **Subsections (Level 2):**
            - Should use decimal numbering: <span class="citation-example">1.1 Background</span>
            
            **Sub-subsections (Level 3):**
            - Should use multi-level decimal numbering: <span class="citation-example">1.1.1 Historical Context</span>
            """, unsafe_allow_html=True)
        
        with tab9:
            st.subheader("Academic Quality Analysis")
            
            # Create table
            table_data = []
            for item in comparison_results["academic_quality"]:
                if item["status"] == "compliant":
                    status_symbol = "✅"
                    status_class = "compliance-compliant"
                elif item["status"] == "partially-compliant":
                    status_symbol = "⚠️"
                    status_class = "compliance-partially-compliant"
                else:
                    status_symbol = "❌"
                    status_class = "compliance-non-compliant"
                
                table_data.append([
                    item["element"], 
                    f'<span class="{status_class}">{status_symbol}</span>', 
                    item["details"]
                ])
            
            # Display table
            st.markdown(
                tabulate(table_data, headers=["Element", "Status", "Details"], tablefmt="html"),
                unsafe_allow_html=True
            )
            
            # Academic quality guidance
            st.markdown("### Academic Writing Quality Guidance")
            st.markdown("""
            **Sentence Length:**
            - Aim for an average sentence length of 15-25 words
            - Vary sentence length for better readability
            - Avoid extremely short (<5 words) or long (>40 words) sentences
            
            **Passive Voice:**
            - Use passive voice sparingly (less than 20% of sentences)
            - Use active voice for clarity and directness
            - Passive voice is appropriate when the actor is unknown or unimportant
            
            **Paragraph Structure:**
            - Each paragraph should focus on a single idea
            - Aim for paragraphs of 3-7 sentences (50-200 words)
            - Use topic sentences to introduce the main idea of each paragraph
            """, unsafe_allow_html=True)
        
        with tab10:
            st.subheader("Journal Citation Analysis")
            
            # Create table
            table_data = []
            for item in comparison_results["journal_citation"]:
                if item["status"] == "compliant":
                    status_symbol = "✅"
                    status_class = "compliance-compliant"
                elif item["status"] == "partially-compliant":
                    status_symbol = "⚠️"
                    status_class = "compliance-partially-compliant"
                else:
                    status_symbol = "❌"
                    status_class = "compliance-non-compliant"
                
                table_data.append([
                    item["element"], 
                    f'<span class="{status_class}">{status_symbol}</span>', 
                    item["details"]
                ])
            
            # Display table
            st.markdown(
                tabulate(table_data, headers=["Element", "Status", "Details"], tablefmt="html"),
                unsafe_allow_html=True
            )
            
            # Journal citation guidance
            st.markdown("### Journal Citation Guidance")
            st.markdown("""
            **QAJ Citation Requirement:**
            
            QAJ requires that all submitted articles cite at least one relevant article from Qubahan Academic Journal. This is a critical requirement for submission acceptance.
            
            **How to Cite QAJ:**
            
            <span class="citation-example">[n] Author A, Author B. Title of article. Qubahan Academic Journal. Year;Volume(Issue):Page range.</span>
            
            You can find relevant QAJ articles to cite at the [Qubahan Academic Journal website](https://journals.qub.edu.iq/).
            """, unsafe_allow_html=True)
        
        with tab11:
            st.subheader("Self-Citation Analysis")
            
            # Create table
            table_data = []
            for item in comparison_results["self_citation"]:
                if item["status"] == "compliant":
                    status_symbol = "✅"
                    status_class = "compliance-compliant"
                elif item["status"] == "partially-compliant":
                    status_symbol = "⚠️"
                    status_class = "compliance-partially-compliant"
                else:
                    status_symbol = "❌"
                    status_class = "compliance-non-compliant"
                
                table_data.append([
                    item["element"], 
                    f'<span class="{status_class}">{status_symbol}</span>', 
                    item["details"]
                ])
            
            # Display table
            st.markdown(
                tabulate(table_data, headers=["Element", "Status", "Details"], tablefmt="html"),
                unsafe_allow_html=True
            )
            
            # Self-citation guidance
            st.markdown("### Self-Citation Guidance")
            st.markdown("""
            **Self-Citation Policy:**
            
            QAJ prohibits self-citation in submitted articles. Self-citation occurs when authors cite their own previous works in the references.
            
            **Why Self-Citation is Prohibited:**
            
            - Maintains objectivity in academic discourse
            - Prevents artificial inflation of citation metrics
            - Ensures diverse perspectives in the literature review
            
            **How to Address Self-Citation:**
            
            If you need to reference your previous work, consider:
            - Citing related work by other researchers instead
            - Focusing on the broader field rather than specific papers
            - Consulting with the journal editor for specific cases
            """, unsafe_allow_html=True)
        
        # Recommendations
        st.subheader("Recommendations")
        
        if comparison_results["major_issues"]:
            st.markdown("Based on the comparison, the following changes are recommended:")
            
            for issue in comparison_results["major_issues"]:
                st.markdown(f"1. **{issue['category']}**: Address {issue['issue'].lower()}")
            
            st.markdown("Refer to the QAJ template document for specific formatting and structure requirements.")
        else:
            st.success("The document appears to be compliant with the QAJ template requirements.")

if __name__ == "__main__":
    main()
